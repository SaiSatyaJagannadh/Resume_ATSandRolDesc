"""Text extraction from resume files (.pdf, .docx, .txt/.md)."""

import os
import re
import tempfile
from pathlib import Path

SUPPORTED = (".pdf", ".docx", ".txt", ".md")


def extract_text(path: str | Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        text = _from_pdf(path)
    elif ext == ".docx":
        text = _from_docx(path)
    elif ext in (".txt", ".md"):
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(
            f"Unsupported file type {ext!r}. Supported types: {', '.join(SUPPORTED)}"
        )

    # Resumes hyperlink the words "LinkedIn" / "GitHub" / "Portfolio" rather than
    # printing the URL. Text extraction keeps the anchor text and drops the href,
    # so the parsed contact block ends up holding the literal string "LinkedIn".
    # The URLs are the only machine-followable thing on the page — recover them.
    links = _pdf_links(path) if ext == ".pdf" else _docx_links(path) if ext == ".docx" else []
    return _normalize(_append_links(text, links))


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Streamlit hands uploads over as bytes, so round-trip through a temp file."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(
            f"Unsupported file type {suffix!r}. Supported types: {', '.join(SUPPORTED)}"
        )
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    try:
        tmp.write(data)
        tmp.close()
        return extract_text(tmp.name)
    finally:
        os.unlink(tmp.name)


def _from_pdf(path: Path) -> str:
    # pdfplumber has better layout handling, but chokes on some malformed PDFs;
    # pypdf is more permissive, so it gets a second shot before we give up.
    text = ""
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        pass

    if not text.strip():
        try:
            from pypdf import PdfReader

            text = "\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
        except Exception:
            text = ""

    if not text.strip():
        # Returning "" here would let the LLM confidently parse an empty resume.
        raise ValueError(
            f"No text could be extracted from {path.name}. It appears to be an "
            f"image-only/scanned PDF and needs OCR before it can be parsed."
        )
    return text


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # Many real resumes lay content out in tables; body order isn't exposed by
    # python-docx without XML walking, so tables land after the paragraphs.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts)


def _pdf_links(path: Path) -> list[str]:
    """URLs behind a PDF's link annotations, in page order.

    Best-effort by design: a malformed annotation dict must never break the
    upload, which is otherwise the only hard-fail path in this module.
    """
    urls = []
    try:
        from pypdf import PdfReader

        for page in PdfReader(path).pages:
            for ref in page.get("/Annots") or []:
                try:
                    action = ref.get_object().get("/A") or {}
                except Exception:
                    continue
                uri = action.get("/URI")
                if uri:
                    urls.append(str(uri))
    except Exception:
        return []
    return urls


def _docx_links(path: Path) -> list[str]:
    """External relationship targets from a .docx — its hyperlink table.

    Read straight from the zip: python-docx exposes relationships only through
    the part API, and this is four lines with no extra dependency.
    """
    try:
        import zipfile

        with zipfile.ZipFile(path) as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8", "replace")
        return re.findall(r'Target="(https?://[^"]+)"', rels)
    except Exception:
        return []


def _append_links(text: str, urls: list[str]) -> str:
    """Append a Links: block for URLs the visible text doesn't already show."""
    seen, new = set(), []
    for url in urls:
        url = url.strip()
        # mailto: and friends carry no information the text block lacks.
        if not url.lower().startswith(("http://", "https://")):
            continue
        if url in seen or url in text:
            continue
        seen.add(url)
        new.append(url)
    if not new:
        return text
    return text + "\n\nLinks:\n" + "\n".join(new)


def _normalize(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
