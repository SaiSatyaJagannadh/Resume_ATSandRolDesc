"""Render a ParsedResume to an ATS-parseable .docx.

ATS parsers read the document as a linear text stream: tables, text boxes,
images, columns and real headers/footers all get mangled or dropped. So this
renderer emits nothing but single-column paragraphs in a standard font, under
the exact section headings ATS section-detection looks for.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

from graph.state import ParsedResume


def render_docx(resume: ParsedResume, out_path: str | Path) -> Path:
    out_path = Path(out_path).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    c = resume.contact
    if c.name:
        doc.add_paragraph(c.name).runs[0].bold = True
    contact_line = " | ".join(v for v in (c.email, c.phone, c.location) if v)
    if contact_line:
        doc.add_paragraph(contact_line)
    links = " | ".join(v for v in (c.linkedin, c.website) if v)
    if links:
        doc.add_paragraph(links)

    if resume.summary:
        _heading(doc, "SUMMARY")
        doc.add_paragraph(resume.summary)

    if resume.skills:
        _heading(doc, "SKILLS")
        doc.add_paragraph(", ".join(resume.skills))

    if resume.experience:
        _heading(doc, "EXPERIENCE")
        for e in resume.experience:
            doc.add_paragraph(f"{e.title} | {e.company}").runs[0].bold = True
            meta = " | ".join(v for v in (e.dates, e.location) if v)
            if meta:
                doc.add_paragraph(meta)
            for b in e.bullets:
                _bullet(doc, b)

    if resume.education:
        _heading(doc, "EDUCATION")
        for ed in resume.education:
            doc.add_paragraph(f"{ed.degree} | {ed.institution}").runs[0].bold = True
            for extra in (ed.dates, ed.details):
                if extra:
                    doc.add_paragraph(extra)

    if resume.projects:
        _heading(doc, "PROJECTS")
        for p in resume.projects:
            doc.add_paragraph(p.name).runs[0].bold = True
            if p.description:
                doc.add_paragraph(p.description)
            for b in p.bullets:
                _bullet(doc, b)

    if resume.certifications:
        _heading(doc, "CERTIFICATIONS")
        for cert in resume.certifications:
            _bullet(doc, cert)

    doc.save(str(out_path))
    return out_path


def _heading(doc: Document, text: str) -> None:
    # Plain bold paragraph, not add_heading: heading styles are fine but the
    # literal uppercase word is what ATS section-detection keys on.
    doc.add_paragraph("")
    doc.add_paragraph(text).runs[0].bold = True


def _bullet(doc: Document, text: str) -> None:
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph(f"• {text}")
