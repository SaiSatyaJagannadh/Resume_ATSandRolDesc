import sys
from pathlib import Path

import pytest
from docx import Document

# No conftest/packaging in this repo, so put the repo root on sys.path.
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from graph.state import (
    Contact,
    EducationEntry,
    ExperienceEntry,
    ParsedResume,
    ProjectEntry,
)
from tools.extract_text import extract_text, extract_text_from_bytes
from tools.render_docx import render_docx


def _resume(**overrides) -> ParsedResume:
    data = dict(
        contact=Contact(name="Ada Lovelace", email="ada@example.com", phone="555-0100"),
        summary="Analytical engine specialist.",
        skills=["Python", "Algorithms"],
        experience=[
            ExperienceEntry(
                company="Analytical Engine Co",
                title="Lead Programmer",
                dates="1842 - 1843",
                bullets=["Wrote the first published algorithm for a machine."],
            )
        ],
        education=[EducationEntry(institution="Home Tutelage", degree="Mathematics")],
        projects=[ProjectEntry(name="Note G", bullets=["Bernoulli number routine."])],
        certifications=["Certified Difference Engineer"],
    )
    data.update(overrides)
    return ParsedResume(**data)


def test_round_trip(tmp_path):
    out = render_docx(_resume(), tmp_path / "r.docx")
    text = extract_text(out)
    assert "Analytical Engine Co" in text
    assert "Wrote the first published algorithm for a machine." in text
    assert "ada@example.com" in text
    for heading in ("SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION", "PROJECTS", "CERTIFICATIONS"):
        assert heading in text


def test_empty_sections_omitted(tmp_path):
    out = render_docx(_resume(projects=[], certifications=[], summary=""), tmp_path / "r.docx")
    text = extract_text(out)
    assert "PROJECTS" not in text
    assert "CERTIFICATIONS" not in text
    assert "SUMMARY" not in text
    assert "EXPERIENCE" in text


def test_no_tables_or_images(tmp_path):
    out = render_docx(_resume(), tmp_path / "r.docx")
    doc = Document(str(out))
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 0


def test_unsupported_extension(tmp_path):
    bad = tmp_path / "resume.rtf"
    bad.write_text("nope")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(bad)


def test_extract_from_bytes_docx(tmp_path):
    out = render_docx(_resume(), tmp_path / "r.docx")
    assert "ada@example.com" in extract_text_from_bytes(out.read_bytes(), "upload.docx")


def test_docx_table_cells_extracted(tmp_path):
    doc = Document()
    doc.add_paragraph("Header paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = "Fortran"
    path = tmp_path / "tabled.docx"
    doc.save(str(path))

    text = extract_text(path)
    assert "Header paragraph" in text
    assert "Skills" in text
    assert "Fortran" in text


def test_repair_skills_splits_unbalanced_brackets():
    from graph.nodes.resume_parser import repair_skills

    # The parser splits dense skill lists on commas, including inside brackets.
    assert repair_skills(["AWS (S3", "EC2", "Redshift)"]) == ["AWS", "S3", "EC2", "Redshift"]
    # Balanced entries survive intact.
    assert repair_skills(["Azure Data Factory (ADF)"]) == ["Azure Data Factory (ADF)"]
    # Empties and dupes are dropped.
    assert repair_skills(["Python", " ", "Python", "BI Tools (Tableau"]) == [
        "Python", "BI Tools", "Tableau"
    ]
